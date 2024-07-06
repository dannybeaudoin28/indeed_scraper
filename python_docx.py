from docx import Document

def get_old_resume_info():
    document = Document("assets\Beaudoin, Danny .docx")
    header = document.sections[0].header
    header_text = '\n'.join([paragraph.text for paragraph in header.paragraphs])

    old_resume = ""

    for text in document.paragraphs:
        old_resume += text.text
            
    return (old_resume, header_text)

def create_resume(content, file_name, old_header, dir_name):
    save_path = "assets\\" + dir_name + "\\Beaudoin, Danny - " + file_name + ".docx"
     
    document = Document()
    section = document.sections[0]
    header = section.header
    
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.text = old_header
    
    document.add_paragraph(content)
    document.save(save_path)